import os
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


OUTPUT_FOLDER = "static/outputs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def save_as_docx(text: str, summary: str, filename: str) -> str:
    path = os.path.join(OUTPUT_FOLDER, filename)
    doc = Document()
    doc.add_heading("Biên bản cuộc họp", 0)
    doc.add_heading("Văn bản gốc", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Tóm tắt", level=1)
    doc.add_paragraph(summary)
    doc.save(path)
    return os.path.join("outputs", filename)


def save_as_pdf(text: str, summary: str, filename: str) -> str:
    path = os.path.join(OUTPUT_FOLDER, filename)
    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Biên bản cuộc họp", styles["Title"]))
    story.append(Paragraph("Văn bản gốc", styles["Heading1"]))
    story.append(Paragraph(text, styles["Normal"]))
    story.append(Paragraph("Tóm tắt", styles["Heading1"]))
    story.append(Paragraph(summary, styles["Normal"]))
    doc.build(story)
    return os.path.join("outputs", filename)
